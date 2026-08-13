"""Cover page, table of contents, headers and footers."""
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0x2E, 0x86, 0xC1)
GREY = RGBColor(0x5D, 0x6D, 0x7E)


def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MIGRATION VALIDATOR")
    r.font.size = Pt(40)
    r.font.bold = True
    r.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Technical Documentation & Handover Guide")
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT

    # divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), "2E86C1")
    pbdr.append(bottom)
    pPr.append(pbdr)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Data Migration Validation Framework")
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PostgreSQL  •  MSSQL  •  Athena  →  Snowflake")
    r.font.size = Pt(12)
    r.font.color.rgb = PRIMARY

    for _ in range(6):
        doc.add_paragraph()

    for line, sz, bold in [
        ("Version 1.0", 12, True),
        ("Comprehensive Reference for New Team Members", 11, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(sz)
        r.font.bold = bold
        r.font.color.rgb = GREY
    doc.add_page_break()


def build_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run("Table of Contents")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = PRIMARY
    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose 'Update Field' to build the Table of Contents."
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)
    doc.add_page_break()


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Migration Validator — Technical Documentation      |      Page ")
        r.font.size = Pt(8)
        r.font.color.rgb = GREY
        # page number field
        fldSimple = OxmlElement("w:fldSimple")
        fldSimple.set(qn("w:instr"), "PAGE")
        p._p.append(fldSimple)
