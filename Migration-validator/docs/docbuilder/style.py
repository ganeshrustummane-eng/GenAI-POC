"""
Shared styling helpers for building the Migration Validator Word documentation.
Provides: document setup, styled headings, justified body text, bullet lists,
code blocks, tables, callout boxes, and image insertion.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand palette ----
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0x2E, 0x86, 0xC1)
DARK = RGBColor(0x21, 0x25, 0x29)
GREY = RGBColor(0x5D, 0x6D, 0x7E)
CODE_BG = "F2F4F6"
CALLOUT_BG = "EBF5FB"
HDR_BG = "1F4E79"


def new_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(8)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    return doc


def _shade(element, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    element.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = PRIMARY
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "2E86C1")
        pbdr.append(bottom)
        pPr.append(pbdr)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = ACCENT
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = GREY
    return p


def add_body(doc, text, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    # Support **bold** inline segments
    parts = text.split("**")
    for i, seg in enumerate(parts):
        run = p.add_run(seg)
        if i % 2 == 1:
            run.font.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    parts = text.split("**")
    for i, seg in enumerate(parts):
        run = p.add_run(seg)
        if i % 2 == 1:
            run.font.bold = True
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def add_code(doc, code):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _shade(cell._tc.get_or_add_tcPr(), CODE_BG)
    cell.paragraphs[0].text = ""
    for i, line in enumerate(code.rstrip("\n").split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK
    _set_borders(tbl, "D0D5DB")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def add_callout(doc, title, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _shade(cell._tc.get_or_add_tcPr(), CALLOUT_BG)
    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    r.font.bold = True
    r.font.color.rgb = PRIMARY
    r2 = p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_borders(tbl, "2E86C1")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i]._tc.get_or_add_tcPr(), HDR_BG)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
    for r_i, row in enumerate(rows):
        cells = tbl.add_row().cells
        for c_i, val in enumerate(row):
            if r_i % 2 == 1:
                _shade(cells[c_i]._tc.get_or_add_tcPr(), "F2F4F6")
            p = cells[c_i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    return tbl


def add_image(doc, path, caption=None, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.add_run(caption)
        run.font.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GREY
        c.paragraph_format.space_after = Pt(10)


def _set_borders(tbl, color):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def page_break(doc):
    doc.add_page_break()
